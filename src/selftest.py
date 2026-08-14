"""빌드된 exe 안에서 분할 로직/라이브러리 번들링이 제대로 되었는지 확인하는 콘솔 자체 테스트."""

import os
import sys
import tempfile
import traceback

import splitter
from docx import Document
from hwpx.document import HwpxDocument

from txt_handler import split_txt_file
from docx_handler import split_docx_file
from hwpx_handler import split_hwpx_file

# import josa  # 이름·호칭 기능 비활성화
# import rename_apply  # 이름·호칭 기능 비활성화
# from term_dict import TermDict  # 이름·호칭 기능 비활성화
import merge_apply
import convert_apply
import batch_rename_apply
import update_checker
import tkinter as tk
import fonts
import prefs


def check_no_mid_sentence_cut(chunks_text, full_text):
    ends = set(splitter.find_sentence_ends(full_text))
    ends.add(len(full_text))
    joined = "".join(chunks_text)
    if joined != full_text:
        raise AssertionError("청크를 이어붙인 결과가 원문과 다릅니다.")
    pos = 0
    for c in chunks_text:
        pos += len(c)
        if pos not in ends:
            raise AssertionError(f"문장 중간에서 잘림: 위치 {pos}")


def main():
    tmp = tempfile.mkdtemp(prefix="tsplit_selftest_")

    text = (
        "이것은 첫 문장입니다. 이것은 두번째 문장이고 조금 더 깁니다. "
        "\"인용문도 있습니다.\" 마지막 문장입니다!"
    )

    txt_path = os.path.join(tmp, "s.txt")
    with open(txt_path, "w", encoding="cp949") as f:
        f.write(text)
    out_paths, enc = split_txt_file(txt_path, 15, tmp)
    joined = ""
    for p in out_paths:
        with open(p, encoding="utf-8-sig") as f:
            joined += f.read()
    assert joined == text, "txt 왕복 결과가 원문과 다릅니다"
    print(f"[OK] txt (encoding={enc}, chunks={len(out_paths)})")

    docx_path = os.path.join(tmp, "s.docx")
    doc = Document()
    doc.add_paragraph(text)
    doc.save(docx_path)
    full_text_docx = "".join(p.text for p in Document(docx_path).paragraphs)
    out_paths = split_docx_file(docx_path, 15, tmp)
    chunk_texts = ["".join(p.text for p in Document(p2).paragraphs) for p2 in out_paths]
    check_no_mid_sentence_cut(chunk_texts, full_text_docx)
    print(f"[OK] docx (chunks={len(out_paths)})")

    hwpx_path = os.path.join(tmp, "s.hwpx")
    hdoc = HwpxDocument.new()
    hdoc.add_paragraph(text)
    for p in list(hdoc.paragraphs):
        if p.text == "":
            p.remove()
    hdoc.save_to_path(hwpx_path)
    full_text_hwpx = "".join(p.text for p in HwpxDocument.open(hwpx_path).paragraphs)
    out_paths = split_hwpx_file(hwpx_path, 15, tmp)
    chunk_texts = ["".join(p.text for p in HwpxDocument.open(p2).paragraphs) for p2 in out_paths]
    check_no_mid_sentence_cut(chunk_texts, full_text_hwpx)
    print(f"[OK] hwpx (chunks={len(out_paths)})")

    # 이름·호칭 기능 비활성화 - josa/term_dict/rename_apply 테스트도 함께 뺐다.
    # assert josa.correct_particle("은", "언니") == "는"
    # assert josa.correct_particle("는", "형") == "은"
    # assert josa.correct_particle("으로", "달") == "로"
    #
    # dict_path = os.path.join(tmp, "terms.csv")
    # td = TermDict(path=dict_path)
    # assert len(td.entries) > 50, "기본 호칭어 시딩 실패"
    # td.add("형", "언니")
    # mapping = td.active_mapping([])
    # assert mapping.get("형") == "언니"
    # renamed = rename_apply.apply_to_text("형은 학교에 갔다.", mapping)
    # assert renamed == "언니는 학교에 갔다.", renamed
    # print("[OK] josa/term_dict/rename_apply")

    m1 = os.path.join(tmp, "m1.txt")
    m2 = os.path.join(tmp, "m2.txt")
    with open(m1, "w", encoding="utf-8") as f:
        f.write("첫 번째 내용")
    with open(m2, "w", encoding="utf-8") as f:
        f.write("두 번째 내용")
    merged_path = os.path.join(tmp, "merged.txt")
    merge_apply.merge_files(
        [
            {"path": m1, "title": "프롤로그"},
            {"path": m2, "title": ""},
        ],
        merged_path,
    )
    with open(merged_path, encoding="utf-8-sig") as f:
        merged_content = f.read()
    assert merged_content == "프롤로그\n\n첫 번째 내용\n\n두 번째 내용", merged_content
    print("[OK] merge_apply")

    segments, truncated = merge_apply.build_preview_segments(
        [{"path": m1, "title": "프롤로그"}, {"path": m2, "title": ""}],
        blank_lines=0,
        header_bold=True,
    )
    assert not truncated
    assert "".join(s[0] for s in segments) == "프롤로그\n\n첫 번째 내용\n두 번째 내용"
    print("[OK] merge_apply preview")

    docx_conv = convert_apply.convert_file(m1, ".docx", tmp)
    assert os.path.isfile(docx_conv)
    txt_back = convert_apply.convert_file(docx_conv, ".txt", tmp)
    with open(txt_back, encoding="utf-8-sig") as f:
        assert f.read() == "첫 번째 내용"
    hwpx_conv = convert_apply.convert_file(docx_conv, ".hwpx", tmp)
    assert os.path.isfile(hwpx_conv)
    docx_back = convert_apply.convert_file(hwpx_conv, ".docx", tmp)
    assert "".join(p.text for p in Document(docx_back).paragraphs) == "첫 번째 내용"
    print("[OK] convert_apply (docx<->hwpx 포함)")

    # 분할/병합 결과를 원본과 다른 확장자로 저장하는 기능 (merge_apply.merge_files의
    # output_path 확장자 분기, split_page._convert_outputs와 같은 절차).
    merged_hwpx = os.path.join(tmp, "merged_as_hwpx.hwpx")
    merge_apply.merge_files([{"path": m1, "title": "프롤로그"}, {"path": m2, "title": ""}], merged_hwpx)
    assert os.path.isfile(merged_hwpx)
    merged_text = "\n".join(p.text for p in HwpxDocument.open(merged_hwpx).paragraphs)
    assert "첫 번째 내용" in merged_text and "두 번째 내용" in merged_text and "프롤로그" in merged_text
    print("[OK] merge_apply (txt 입력 -> hwpx로 저장)")

    long_text = "이것은 문장입니다. " * 200
    hdoc = HwpxDocument.new()
    default_paragraphs = list(hdoc.paragraphs)
    hdoc.add_paragraph(long_text)
    for p in default_paragraphs:
        p.remove()
    split_src_hwpx = os.path.join(tmp, "split_src.hwpx")
    hdoc.save_to_path(split_src_hwpx)
    native_parts = split_hwpx_file(split_src_hwpx, 500, tmp)
    assert len(native_parts) > 1
    txt_parts = []
    for p in native_parts:
        converted = convert_apply.convert_file(p, ".txt", tmp)
        os.remove(p)
        txt_parts.append(converted)
    for p in txt_parts:
        with open(p, encoding="utf-8-sig") as f:
            assert f.read().strip(), f"empty part: {p}"
    print("[OK] hwpx 분할 결과를 txt로 저장")

    r1 = os.path.join(tmp, "r1.txt")
    with open(r1, "w", encoding="utf-8") as f:
        f.write("x")
    succeeded, failed, skipped = batch_rename_apply.apply_renames([{"path": r1, "new_base": "renamed"}])
    assert succeeded and os.path.isfile(os.path.join(tmp, "renamed.txt"))
    print("[OK] batch_rename_apply")

    assert update_checker.is_newer("v1.2.0", "1.1.0") is True
    assert update_checker.is_newer("1.0.0", "1.0.0") is False
    assert update_checker.fetch_latest_release(repo="this-should-not-exist-xyz/no-repo") is None
    print("[OK] update_checker")

    previews, ptruncated, ptotal = splitter.build_boundary_previews(text, 15)
    assert ptotal >= 2
    assert not ptruncated
    assert previews[0]["is_last"] is False
    print("[OK] splitter.build_boundary_previews")

    root = tk.Tk()
    root.withdraw()
    available = fonts.list_available_fonts()
    assert len(available) > 10
    assert fonts.resolve_default_font("ko", available) == "맑은 고딕"
    fonts.apply_font_family("Consolas")
    assert fonts.current_family() == "Consolas"
    prefs_path = os.path.join(tmp, "prefs.json")
    prefs.set_pref("font_family", "Consolas", path=prefs_path)
    assert prefs.load_prefs(prefs_path).get("font_family") == "Consolas"
    root.destroy()
    print("[OK] fonts/prefs")

    print("ALL SELFTEST PASSED")


if __name__ == "__main__":
    try:
        main()
        sys.exit(0)
    except Exception:
        traceback.print_exc()
        sys.exit(1)

"""여러 파일(txt/docx/hwpx)을 순서대로 하나로 합치는 로직.

entries는 다음 형태의 dict 리스트다.
    {
        "path": "C:/.../1화.txt",
        "title": "1화 제목",   # 비워두면 이 파일 앞에 제목을 넣지 않는다.
    }
순서는 리스트의 순서를 그대로 따른다.

blank_lines: 파일과 파일 사이에 넣을 빈 줄 수 (0 = 줄바꿈만, 1 = 빈 줄 하나, ...).
header_bold / header_italic: 제목 줄에 적용할 서식
(txt에는 서식 개념이 없으므로 무시된다).
"""

import copy
import os

from docx import Document
from hwpx.document import HwpxDocument

from txt_handler import read_text_auto

DEFAULT_BLANK_LINES = 1


class MixedExtensionError(Exception):
    def __init__(self, exts):
        self.exts = sorted(exts)
        super().__init__(f"mixed extensions: {self.exts}")


def default_title(path):
    return os.path.splitext(os.path.basename(path))[0]


def header_lines(entry):
    """entry에 제목이 있으면 [제목]을, 비어 있으면 빈 리스트를 반환(헤더 없음)."""
    title = (entry.get("title") or "").strip()
    return [title] if title else []


def common_extension(entries):
    if not entries:
        raise ValueError("합칠 파일이 없습니다.")
    exts = {os.path.splitext(e["path"])[1].lower() for e in entries}
    if len(exts) > 1:
        raise MixedExtensionError(exts)
    return next(iter(exts))


def read_plain_text(path):
    """미리보기·txt 합치기에서 쓰는, 파일 하나의 순수 텍스트(문단은 줄바꿈으로 구분)."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        text, _encoding = read_text_auto(path)
        return text
    if ext == ".docx":
        return "\n".join(p.text for p in Document(path).paragraphs)
    if ext == ".hwpx":
        return "\n".join(p.text for p in HwpxDocument.open(path).paragraphs)
    raise ValueError(f"지원하지 않는 형식입니다: {ext}")


# ---------- txt ----------

def _txt_block(entry):
    lines = header_lines(entry)
    text = read_plain_text(entry["path"])
    if lines:
        return "\n".join(lines) + "\n\n" + text
    return text


def merge_txt_files(entries, output_path, blank_lines=DEFAULT_BLANK_LINES):
    blocks = [_txt_block(e) for e in entries]
    sep = "\n" * (blank_lines + 1)
    content = sep.join(blocks)
    with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
        f.write(content)


# ---------- docx ----------
# 문단을 이루는 XML(<w:p>)을 그대로 복사해서 붙여 넣기 때문에, 문단/글자 서식(굵게·기울임·
# 글꼴 등 직접 지정된 서식)은 모든 파일에 대해 그대로 유지된다. 표/이미지는 대상이 아니다.

def _sectpr_index(body):
    for i, child in enumerate(body):
        if child.tag.endswith("}sectPr"):
            return i
    return len(body)


def _append_docx_header(dest_doc, text, bold=True, italic=False):
    p = dest_doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def _append_docx_source(dest_doc, src_path):
    src_doc = Document(src_path)
    body = dest_doc.element.body
    insert_at = _sectpr_index(body)
    for p in src_doc.paragraphs:
        new_p = copy.deepcopy(p._p)
        body.insert(insert_at, new_p)
        insert_at += 1


def merge_docx_files(entries, output_path, blank_lines=DEFAULT_BLANK_LINES, header_bold=True, header_italic=False):
    dest = Document()
    for i, e in enumerate(entries):
        if i > 0:
            for _ in range(blank_lines):
                dest.add_paragraph("")
        for line in header_lines(e):
            _append_docx_header(dest, line, bold=header_bold, italic=header_italic)
        _append_docx_source(dest, e["path"])
    dest.save(output_path)


# ---------- hwpx ----------
# 첫 번째 파일을 기준 문서로 열어 그 서식(스타일 정의)을 그대로 사용한다.
# 첫 번째 파일 본문은 원래 서식이 그대로 유지되지만, 이후 파일들은 서로 다른 문서에서
# 정의된 스타일 번호가 충돌할 수 있어 안전하게 일반 텍스트로만 합쳐진다.
# (제목/에피소드/챕터 같은 헤더 줄은 대상 문서 자체의 서식을 새로 적용하므로 모든
# 파일에 대해 정상적으로 굵게/기울임이 표시된다.)

def _append_hwpx_header(dest_doc, text, anchor=None, bold=True, italic=False):
    p = dest_doc.add_paragraph(text)
    if p.runs:
        p.runs[0].bold = bold
        p.runs[0].italic = italic
    if anchor is not None:
        anchor.element.addprevious(p.element)
    return p


def merge_hwpx_files(entries, output_path, blank_lines=DEFAULT_BLANK_LINES, header_bold=True, header_italic=False):
    first_path = entries[0]["path"]
    dest = HwpxDocument.open(first_path)

    existing_paragraphs = list(dest.paragraphs)
    anchor = existing_paragraphs[0] if existing_paragraphs else None
    for line in header_lines(entries[0]):
        _append_hwpx_header(dest, line, anchor=anchor, bold=header_bold, italic=header_italic)

    for e in entries[1:]:
        for _ in range(blank_lines):
            dest.add_paragraph("")
        for line in header_lines(e):
            _append_hwpx_header(dest, line, bold=header_bold, italic=header_italic)
        src = HwpxDocument.open(e["path"])
        for sp in src.paragraphs:
            dest.add_paragraph(sp.text)

    dest.save_to_path(output_path)


def merge_files(entries, output_path, blank_lines=DEFAULT_BLANK_LINES, header_bold=True, header_italic=False):
    ext = common_extension(entries)
    if ext == ".txt":
        merge_txt_files(entries, output_path, blank_lines=blank_lines)
    elif ext == ".docx":
        merge_docx_files(entries, output_path, blank_lines=blank_lines, header_bold=header_bold, header_italic=header_italic)
    elif ext == ".hwpx":
        merge_hwpx_files(entries, output_path, blank_lines=blank_lines, header_bold=header_bold, header_italic=header_italic)
    else:
        raise ValueError(f"지원하지 않는 형식입니다: {ext}")


# ---------- 미리보기 ----------

def build_preview_segments(entries, blank_lines=DEFAULT_BLANK_LINES, header_bold=True, header_italic=False, max_chars=50000):
    """[(text, is_header, bold, italic), ...] 세그먼트 리스트와 잘림 여부(bool)를 반환.

    실제 병합 결과와 같은 순서·간격·헤더 서식을 미리 보여주기 위한 것으로, 파일 내용은
    문단을 줄바꿈으로 이어 붙인 일반 텍스트로만 표시한다(글자 단위 서식은 반영하지 않음).
    """
    segments = []
    total = 0
    truncated = False

    for i, entry in enumerate(entries):
        if i > 0:
            segments.append(("\n" * (blank_lines + 1), False, False, False))

        lines = header_lines(entry)
        if lines:
            header_text = "\n".join(lines)
            segments.append((header_text, True, header_bold, header_italic))
            segments.append(("\n\n", False, False, False))
            total += len(header_text)

        try:
            body_text = read_plain_text(entry["path"])
        except Exception as e:
            body_text = f"[읽기 실패: {e}]"

        if total + len(body_text) > max_chars:
            body_text = body_text[: max(0, max_chars - total)]
            truncated = True
            segments.append((body_text, False, False, False))
            break

        segments.append((body_text, False, False, False))
        total += len(body_text)

    return segments, truncated

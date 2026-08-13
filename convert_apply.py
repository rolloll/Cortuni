"""텍스트(.txt) <-> Word(.docx)/한글(.hwpx) 확장자(형식) 변환.

docx/hwpx -> txt는 본문 텍스트만 추출한다(표/이미지/서식 유실).
txt -> docx/hwpx는 줄바꿈을 기준으로 각 줄을 하나의 단락으로 만든다(서식 없음).
docx <-> hwpx는 단락과 글자 단위 굵게/기울임 서식은 유지하되, 표/이미지·글꼴/색상
같은 세부 서식은 변환 대상이 아니다(새 문서를 만들어 다시 구성하는 방식이라 원본
문서의 스타일 정의 자체를 그대로 옮기지는 않는다).
"""

import os

from docx import Document
from hwpx.document import HwpxDocument

from txt_handler import read_text_auto
from merge_apply import read_plain_text


def convert_to_txt(src_path, output_path):
    text = read_plain_text(src_path)
    with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
        f.write(text)


def _split_lines(text):
    """윈도우(\\r\\n)/구형 맥(\\r) 줄바꿈도 모두 \\n 기준 한 줄로 인식하도록 정규화 후 분리."""
    return text.replace("\r\n", "\n").replace("\r", "\n").split("\n")


def convert_txt_to_docx(src_path, output_path):
    text, _encoding = read_text_auto(src_path)
    doc = Document()
    for line in _split_lines(text):
        doc.add_paragraph(line)
    doc.save(output_path)


def convert_txt_to_hwpx(src_path, output_path):
    text, _encoding = read_text_auto(src_path)
    doc = HwpxDocument.new()
    default_paragraphs = list(doc.paragraphs)
    for line in _split_lines(text):
        doc.add_paragraph(line)
    for p in default_paragraphs:
        p.remove()
    doc.save_to_path(output_path)


def convert_docx_to_hwpx(src_path, output_path):
    src_doc = Document(src_path)
    dest = HwpxDocument.new()
    default_paragraphs = list(dest.paragraphs)

    for p in src_doc.paragraphs:
        if not p.text:
            dest.add_paragraph("")
            continue
        new_p = dest.add_paragraph("", include_run=False)
        for run in p.runs:
            if not run.text:
                continue
            new_run = new_p.add_run(run.text)
            new_run.bold = bool(run.bold)
            new_run.italic = bool(run.italic)

    for p in default_paragraphs:
        p.remove()
    dest.save_to_path(output_path)


def convert_hwpx_to_docx(src_path, output_path):
    src_doc = HwpxDocument.open(src_path)
    dest = Document()

    for p in src_doc.paragraphs:
        new_p = dest.add_paragraph()
        for run in p.runs:
            if not run.text:
                continue
            new_run = new_p.add_run(run.text)
            new_run.bold = bool(run.bold)
            new_run.italic = bool(run.italic)

    dest.save(output_path)


CONVERTERS = {
    (".docx", ".txt"): convert_to_txt,
    (".hwpx", ".txt"): convert_to_txt,
    (".txt", ".docx"): convert_txt_to_docx,
    (".txt", ".hwpx"): convert_txt_to_hwpx,
    (".docx", ".hwpx"): convert_docx_to_hwpx,
    (".hwpx", ".docx"): convert_hwpx_to_docx,
}


def supported_targets(src_ext):
    return sorted({dst for (src, dst) in CONVERTERS if src == src_ext})


def convert_file(src_path, target_ext, output_dir):
    src_ext = os.path.splitext(src_path)[1].lower()
    converter = CONVERTERS.get((src_ext, target_ext))
    if converter is None:
        raise ValueError(f"'{src_ext}' -> '{target_ext}' 변환은 지원하지 않습니다.")
    base = os.path.splitext(os.path.basename(src_path))[0]
    output_path = os.path.join(output_dir, f"{base}{target_ext}")
    converter(src_path, output_path)
    return output_path

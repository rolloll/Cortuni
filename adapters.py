"""docx / hwpx 문서의 단락(paragraph)·런(run)을 동일한 인터페이스로 다루기 위한 어댑터.

.text / .runs(런 어댑터 리스트) / .remove()를 갖는 단락 어댑터와
.get_text() / .set_text(str) / .remove()를 갖는 런 어댑터로 통일해서,
분할(doc_split.py)과 이름 바꾸기(rename_apply.py) 양쪽에서 재사용한다.
"""

from docx import Document
from hwpx.document import HwpxDocument


class DocxRunAdapter:
    def __init__(self, run):
        self._run = run

    def get_text(self):
        return self._run.text or ""

    def set_text(self, text):
        self._run.text = text

    def remove(self):
        element = self._run._element
        element.getparent().remove(element)


class DocxParagraphAdapter:
    def __init__(self, paragraph):
        self._paragraph = paragraph

    @property
    def text(self):
        return self._paragraph.text

    @property
    def runs(self):
        return [DocxRunAdapter(r) for r in self._paragraph.runs]

    def remove(self):
        element = self._paragraph._element
        element.getparent().remove(element)


class HwpxRunAdapter:
    def __init__(self, run):
        self._run = run

    def get_text(self):
        return self._run.text or ""

    def set_text(self, text):
        self._run.text = text

    def remove(self):
        self._run.remove()


class HwpxParagraphAdapter:
    def __init__(self, paragraph):
        self._paragraph = paragraph

    @property
    def text(self):
        return self._paragraph.text

    @property
    def runs(self):
        return [HwpxRunAdapter(r) for r in self._paragraph.runs]

    def remove(self):
        self._paragraph.remove()


def open_docx(path):
    return Document(path)


def docx_paragraphs(doc):
    return [DocxParagraphAdapter(p) for p in doc.paragraphs]


def save_docx(doc, path):
    doc.save(path)


def open_hwpx(path):
    return HwpxDocument.open(path)


def hwpx_paragraphs(doc):
    return [HwpxParagraphAdapter(p) for p in doc.paragraphs]


def save_hwpx(doc, path):
    doc.save_to_path(path)
